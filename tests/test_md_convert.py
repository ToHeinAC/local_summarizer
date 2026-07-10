"""Markdown conversion: per-page PDF routing, OCR prompt choice, DOCX mapping."""

import io

import pytest
from docx import Document

from src import md_convert


def _docx(build) -> bytes:
    document = Document()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_text_layer_page_yields_text(pdf_bytes):
    kinds = [kind for kind, _ in md_convert.iter_pdf_pages(pdf_bytes, dpi=72)]
    assert kinds == ["text"]


def test_page_without_text_layer_yields_image(scanned_pdf_bytes):
    pages = list(md_convert.iter_pdf_pages(scanned_pdf_bytes, dpi=72))
    assert [kind for kind, _ in pages] == ["image"]
    assert pages[0][1].size[0] > 0  # a real rasterized PIL image


def test_page_below_text_threshold_is_treated_as_scanned():
    """A page with a text layer shorter than TEXT_THRESHOLD still takes the OCR path."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Page 1")  # 6 chars, well under the 40-char threshold
    kinds = [kind for kind, _ in md_convert.iter_pdf_pages(bytes(pdf.output()), dpi=72)]
    assert kinds == ["image"]


def test_deepseek_gets_grounding_prompt(stub_ollama):
    from PIL import Image

    md_convert.convert_image(Image.new("RGB", (4, 4)), "deepseek-ocr:3b")
    assert stub_ollama.ocr_prompts == [md_convert.OCR_DEEPSEEK_PROMPT]


def test_other_vision_model_gets_system_prompt(stub_ollama):
    from PIL import Image

    md_convert.convert_image(Image.new("RGB", (4, 4)), "llava:7b")
    assert md_convert.OCR_SYSTEM_PROMPT in stub_ollama.ocr_prompts[0]
    assert md_convert.OCR_USER_PROMPT in stub_ollama.ocr_prompts[0]


def test_rewrite_prompt_prepends_source_text(stub_ollama):
    md_convert.rewrite_text("Hello world", "gemma4:e4b")
    assert stub_ollama.rewritten[0].endswith("Hello world")
    assert stub_ollama.rewritten[0].startswith(md_convert.MD_REWRITE_PROMPT[:20])


def test_ocr_model_unloaded_only_when_ocr_ran(scanned_pdf_bytes, pdf_bytes, stub_ollama):
    md_convert.pdf_to_markdown(pdf_bytes, "deepseek-ocr:3b", "gemma4:e4b", 72)
    assert stub_ollama.unloaded == [], "no OCR ran, so nothing to evict"

    md_convert.pdf_to_markdown(scanned_pdf_bytes, "deepseek-ocr:3b", "gemma4:e4b", 72)
    assert stub_ollama.unloaded == ["deepseek-ocr:3b"]


def test_pdf_progress_reports_every_page(pdf_bytes, stub_ollama):
    events = []
    md_convert.pdf_to_markdown(
        pdf_bytes, "deepseek-ocr:3b", "gemma4:e4b", 72,
        on_progress=lambda done, total, label: events.append((done, total, label)),
    )
    assert events[0][2].startswith("Seite 1/1")
    assert events[-1][:2] == (1, 1)


def test_docx_maps_headings_and_lists():
    def build(doc):
        doc.add_heading("Title", level=1)
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Body text")
        doc.add_paragraph("A point", style="List Bullet")

    md = md_convert.docx_to_markdown(_docx(build))
    assert "# Title" in md
    assert "## Section" in md
    assert "- A point" in md
    assert "Body text" in md


def test_docx_maps_tables():
    def build(doc):
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Year"
        table.rows[0].cells[1].text = "Revenue"
        table.rows[1].cells[0].text = "2024"
        table.rows[1].cells[1].text = "4.2M"

    md = md_convert.docx_to_markdown(_docx(build))
    assert "| Year | Revenue |" in md
    assert "| --- | --- |" in md
    assert "| 2024 | 4.2M |" in md


def test_docx_preserves_document_order():
    def build(doc):
        doc.add_paragraph("first")
        doc.add_table(rows=1, cols=1).rows[0].cells[0].text = "middle"
        doc.add_paragraph("last")

    md = md_convert.docx_to_markdown(_docx(build))
    assert md.index("first") < md.index("middle") < md.index("last")


def test_unsupported_suffix_raises_via_extract():
    from src.extract import UnsupportedFileError, to_markdown

    with pytest.raises(UnsupportedFileError):
        to_markdown("x.rtf", b"data")
