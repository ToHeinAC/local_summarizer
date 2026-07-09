"""Render a Markdown summary to md/docx/pdf bytes. Plain python, no LangChain.

Handles a small Markdown subset produced by the templates: ATX headings
(``#``..``###``) and ``-``/``*`` bullet lists; everything else is a paragraph.
"""

from __future__ import annotations

import io
import os
import re

from docx import Document
from fpdf import FPDF

# Optional system Unicode font so non-latin summaries render in the PDF.
_UNICODE_FONT_PATHS = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
)

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")


def _parse_line(line: str) -> tuple[str, int, str]:
    """Classify a Markdown line -> (kind, level, text)."""
    stripped = line.strip()
    if not stripped:
        return ("blank", 0, "")
    heading = _HEADING.match(stripped)
    if heading:
        return ("heading", len(heading.group(1)), heading.group(2).strip())
    bullet = _BULLET.match(stripped)
    if bullet:
        return ("bullet", 0, bullet.group(1).strip())
    return ("paragraph", 0, stripped)


def to_markdown(summary: str) -> bytes:
    """Return the summary as UTF-8 Markdown bytes."""
    return summary.encode("utf-8")


def to_docx(summary: str) -> bytes:
    """Render the summary to a .docx document."""
    document = Document()
    for line in summary.splitlines():
        kind, level, text = _parse_line(line)
        if kind == "heading":
            document.add_heading(text, level=min(level, 4))
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "paragraph":
            document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _find_unicode_font() -> str | None:
    for path in _UNICODE_FONT_PATHS:
        if os.path.exists(path):
            return path
    return None


def to_pdf(summary: str) -> bytes:
    """Render the summary to a PDF using fpdf2 (pure python)."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    font_path = _find_unicode_font()
    if font_path:
        pdf.add_font("body", "", font_path)
        family, unicode_ok = "body", True
    else:
        family, unicode_ok = "Helvetica", False

    def write(text: str, size: int, gap: float) -> None:
        if not unicode_ok:
            text = text.encode("latin-1", "replace").decode("latin-1")
        pdf.set_font(family, size=size)
        pdf.multi_cell(0, size * 0.6, text)
        pdf.ln(gap)

    for line in summary.splitlines():
        kind, level, text = _parse_line(line)
        if kind == "heading":
            write(text, max(16 - 2 * level, 11), 1)
        elif kind == "bullet":
            write(f"• {text}", 11, 0.5)
        elif kind == "paragraph":
            write(text, 11, 1)

    return bytes(pdf.output())


EXPORTERS = {
    "md": ("text/markdown", to_markdown),
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        to_docx,
    ),
    "pdf": ("application/pdf", to_pdf),
}
