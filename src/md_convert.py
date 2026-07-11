"""Convert uploads to Markdown before summarization. Plain python, no LangChain.

Pipeline ported and adapted from ToHeinAC/KB_BS_local-wiki-he (Apache-2.0),
itself derived from ToHeinAC/MD-maker (Apache-2.0):

- PDF: per-page routing — pages with a text layer are LLM-rewritten into clean
  Markdown (wording preserved); scanned/image-only pages are rasterized and
  OCR'd by a vision model.
- DOCX: deterministic heading/list/table mapping (no LLM).

All Ollama access goes through ``ollama_client``; all prompts live in ``prompts``.
"""

from __future__ import annotations

import base64
import io
from collections.abc import Callable, Iterator
from concurrent.futures import FIRST_COMPLETED, ThreadPoolExecutor, wait

from src import ollama_client
from src.i18n import DEFAULT_LANG, t
from src.prompts import (
    MD_REWRITE_PROMPT,
    OCR_DEEPSEEK_PROMPT,
    OCR_SYSTEM_PROMPT,
    OCR_USER_PROMPT,
)

TEXT_THRESHOLD = 40  # chars; below this a PDF page is treated as image-only
MAX_CONVERT_WORKERS = 4  # parallel page LLM calls; align with OLLAMA_NUM_PARALLEL

# Callback signature: on_progress(done: int, total: int, label: str) -> None
ProgressCb = Callable[[int, int, str], None]


def _image_to_base64(pil_image) -> str:
    buf = io.BytesIO()
    pil_image.convert("RGB").save(buf, format="JPEG", quality=85)
    return base64.b64encode(buf.getvalue()).decode()


def iter_pdf_pages(data: bytes, dpi: int) -> Iterator[tuple[str, object]]:
    """Yield ``('text', str)`` for pages with a text layer, else ``('image', PIL.Image)``."""
    import pypdfium2 as pdfium

    scale = dpi / 72
    for page in pdfium.PdfDocument(data):
        text = page.get_textpage().get_text_range().strip()
        if len(text) >= TEXT_THRESHOLD:
            yield "text", text
        else:
            yield "image", page.render(scale=scale).to_pil()


def _pdf_page_count(data: bytes) -> int:
    import pypdfium2 as pdfium

    return len(pdfium.PdfDocument(data))


def rewrite_text(text: str, model_id: str, host: str | None = None) -> str:
    """Reformat already-extracted PDF text as Markdown without altering wording."""
    return ollama_client.rewrite(model_id, MD_REWRITE_PROMPT + text, host)


def convert_image(pil_image, model_id: str, host: str | None = None) -> str:
    """OCR one page image to Markdown via a vision model."""
    if model_id.startswith("deepseek-ocr"):
        prompt = OCR_DEEPSEEK_PROMPT
    else:
        prompt = f"{OCR_SYSTEM_PROMPT}\n\n{OCR_USER_PROMPT}"
    return ollama_client.ocr(model_id, prompt, _image_to_base64(pil_image), host)


def _paragraph_md(para) -> str | None:
    text = para.text.strip()
    if not text:
        return None
    style = para.style.name
    if style.startswith("Heading 1"):
        return f"# {text}"
    if style.startswith("Heading 2"):
        return f"## {text}"
    if style.startswith("Heading 3"):
        return f"### {text}"
    if "Bullet" in style:
        return f"- {text}"
    if "Number" in style:
        return f"1. {text}"
    return text


def _table_md(table) -> str | None:
    rows = table.rows
    if not rows:
        return None
    header = [c.text.strip() for c in rows[0].cells]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
    return "\n".join(lines)


def docx_to_markdown(data: bytes) -> str:
    """Convert a .docx file to Markdown, preserving document order (no LLM step)."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            md = _paragraph_md(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            md = _table_md(Table(child, doc))
        else:
            md = None
        if md:
            parts.append(md)
    return "\n\n".join(parts)


def pdf_to_markdown(
    data: bytes,
    ocr_model: str,
    rewrite_model: str,
    dpi: int,
    host: str | None = None,
    on_progress: ProgressCb | None = None,
    lang: str = DEFAULT_LANG,
    fast: bool = False,
) -> str:
    """Convert a PDF to Markdown, OCR'ing any page without a usable text layer.

    ``lang`` selects the GUI language of the per-page progress labels.

    ``fast=True`` skips the per-page LLM rewrite for pages that already have a
    text layer: their extracted text is used verbatim (byte-exact wording, no
    model call), so a digital PDF converts in near-zero time. Scanned pages
    still require OCR. Trade-off: no LLM-added headings/tables and raw reading
    order for multi-column layouts.

    Pages are converted concurrently (up to ``MAX_CONVERT_WORKERS`` LLM calls in
    flight) since each page is an independent Ollama request and the per-page
    rewrite/OCR dominates conversion time. Rendering stays on the calling thread
    (pypdfium2 is not thread-safe); only the LLM calls run in the pool. Results
    are slotted back by page index so output order is preserved, and progress
    counts completed pages so the bar stays monotonic despite out-of-order
    completion.
    """
    total = _pdf_page_count(data)
    results: list[str] = [""] * total
    used_ocr = False
    done = 0
    pages = enumerate(iter_pdf_pages(data, dpi))

    def _convert(kind, payload) -> str:
        if kind == "text":
            return payload if fast else rewrite_text(payload, rewrite_model, host)
        return convert_image(payload, ocr_model, host)

    try:
        with ThreadPoolExecutor(max_workers=MAX_CONVERT_WORKERS) as pool:
            pending: dict = {}

            def submit_more() -> None:
                nonlocal used_ocr
                while len(pending) < MAX_CONVERT_WORKERS:
                    try:
                        idx, (kind, payload) = next(pages)
                    except StopIteration:
                        return
                    if kind == "image":
                        used_ocr = True
                    pending[pool.submit(_convert, kind, payload)] = (idx, kind)

            submit_more()
            while pending:
                finished, _ = wait(pending, return_when=FIRST_COMPLETED)
                for fut in finished:
                    idx, kind = pending.pop(fut)
                    results[idx] = fut.result()
                    done += 1
                    if on_progress:
                        kind_label = "OCR" if kind == "image" else "Text"
                        label = t("page", lang, done=done, total=total, kind=kind_label)
                        on_progress(done, total, label)
                submit_more()
    finally:
        # Free the vision model so the summarizer's text model gets the full GPU
        # instead of splitting VRAM with a resident OCR model. Skipped when no
        # page was OCR'd, since unloading would otherwise load it just to evict it.
        if used_ocr:
            ollama_client.unload(ocr_model, host)
    if on_progress:
        on_progress(total, total, t("converted", lang))
    return "\n\n".join(results)
